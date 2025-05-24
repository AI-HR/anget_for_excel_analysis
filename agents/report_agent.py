# agents/report_agent.py
import os
import pandas as pd
from typing import Dict, Any, List, Optional, Tuple
from langchain.agents import initialize_agent, Tool
from langchain.agents import AgentType
from langchain_openai import ChatOpenAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from config.config import OPENAI_API_KEY, MODEL_NAME, MODEL_BASE_URL, OUTPUT_DIR, REPORT_TEMPLATE
from agents.business_context_agent import BusinessContextAgent

class ReportAgent:
    """基于LLM的智能报告生成Agent - 根据标准报告结构生成数据分析报告"""
    
    def __init__(self):
        # 初始化LLM
        self.llm = ChatOpenAI(
            temperature=0.3,
            api_key=OPENAI_API_KEY,
            model_name=MODEL_NAME,
            base_url=MODEL_BASE_URL
        )
        
        # 初始化业务场景识别Agent
        self.business_context_agent = BusinessContextAgent()
        
        # 报告输出目录
        self.output_dir = os.path.join(OUTPUT_DIR, "reports")
        os.makedirs(self.output_dir, exist_ok=True)
        
        # 定义报告生成工具
        self.tools = [
            Tool(
                name="生成Word报告",
                func=self._generate_word_report,
                description="生成Word格式的数据分析报告"
            ),
            Tool(
                name="生成摘要",
                func=self._generate_summary,
                description="生成数据分析结果的摘要"
            ),
            Tool(
                name="生成建议",
                func=self._generate_recommendations,
                description="根据数据分析结果生成业务建议"
            ),
            Tool(
                name="生成引言与背景",
                func=self._generate_introduction,
                description="生成报告的引言与背景部分"
            )
        ]
        
        # 初始化Agent
        self.agent = initialize_agent(
            self.tools,
            self.llm,
            agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
            verbose=True,
            handle_parsing_errors=True
        )
    
    def generate_report(self, df: pd.DataFrame, analysis_results: Dict[str, Any], 
                        visualization_results: Dict[str, Any], report_title: str = None,
                        business_context: Optional[Dict[str, Any]] = None,
                        analysis_themes: Optional[List[Dict[str, Any]]] = None,
                        analysis_outline: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """生成数据分析报告"""
        # 数据概览
        data_info = self._get_data_overview(df)
        
        # 如果没有提供业务场景，则识别业务场景
        if not business_context:
            try:
                print("识别数据业务场景...")
                business_context = self.business_context_agent.identify_business_context(df)
                print(f"识别到业务场景: {business_context['context']}，置信度: {business_context['confidence']:.2f}")
                
                # 生成报告标题和副标题
                title_info = self.business_context_agent.generate_report_title(business_context, df)
                if "title" in title_info:
                    report_title = title_info["title"]
                    report_subtitle = title_info.get("subtitle", "")
                else:
                    report_title = f"{business_context['context']}数据分析报告"
                    report_subtitle = "基于专业分析维度的深度洞察"
            except Exception as e:
                print(f"业务场景识别失败: {str(e)}")
                business_context = {"context": "未知", "analysis_dimensions": []}
                report_title = "数据分析报告"
                report_subtitle = "多维度数据分析与洞察"
        else:
            if not report_title:
                report_title = f"{business_context['context']}数据分析报告"
                report_subtitle = "基于专业分析维度的深度洞察"
        
        # 准备业务场景信息
        context_info = ""
        if business_context and "context" in business_context:
            context_info = f"业务场景: {business_context['context']}\n"
            if "confidence" in business_context:
                context_info += f"置信度: {business_context['confidence']:.2f}\n"
            if "explanation" in business_context and business_context["explanation"]:
                context_info += f"解释: {business_context['explanation']}\n\n"
            if "analysis_dimensions" in business_context and business_context["analysis_dimensions"]:
                context_info += "专业分析维度:\n"
                for dimension in business_context['analysis_dimensions']:
                    context_info += f"- {dimension['name']}\n"
                    if "description" in dimension and dimension['description']:
                        context_info += f"  {dimension['description']}\n"
        else:
            context_info = "业务场景识别失败，将进行通用数据分析。\n"
        
        # 准备分析主题信息
        themes_info = ""
        if analysis_themes:
            themes_info = "分析主题:\n"
            for theme in analysis_themes:
                themes_info += f"- {theme['name']}: {theme['description']}\n"
                if "key_questions" in theme and theme["key_questions"]:
                    themes_info += f"  关键问题: {', '.join(theme['key_questions'])}\n"
        
        # 准备分析大纲信息
        outline_info = ""
        if analysis_outline and "outline" in analysis_outline:
            outline_info = "分析大纲结构:\n"
            for section, items in analysis_outline["outline"].items():
                outline_info += f"- {section}: {', '.join(items)}\n"
        
        # 构建Agent提示
        prompt = f"""
        我需要为一个数据分析项目生成完整的报告。数据集信息如下：
        
        {data_info}
        
        业务场景信息：
        {context_info}
        
        {themes_info}
        
        {outline_info}
        
        分析结果概要：
        {self._format_analysis_results(analysis_results)}
        
        可视化图表概要：
        {self._format_visualization_results(visualization_results)}
        
        请根据以上信息，生成一份专业、全面的数据分析报告，严格遵循以下标准报告结构：
        
        1. 报告封面与摘要
           - 标题：体现核心分析主题
           - 副标题：说明分析维度
           - 摘要：概述数据集、业务场景、分析目标和主要发现
           
        2. 引言与背景（200-300字）
           - 分析背景：业务需求来源、行业趋势或市场环境说明
           - 分析目标：待解决的3-5个核心问题和成功指标
           
        3. 数据与方法
           - 数据概况：数据字段、数据质量和时间跨度
           - 分析方法：使用的分析方法及其原因
           
        4. 分析结果呈现
           - 核心指标仪表盘：5-7个KPI看板
           - 专题分析模块：1-3个图表/模块及其解读
           - 关键发现标注：异常数据点和特殊事件说明
           
        5. 结论与建议（300-400字）
           - 核心结论：验证假设的结果和战略性发现
           - 可行性建议：短期行动、长期规划和创新方向
           
        请确保报告专业、清晰，突出最重要的发现，并紧密结合识别出的业务场景进行分析。
        """
        
        # 运行Agent获取报告计划
        try:
            report_plan = self.agent.invoke(prompt)
            # 执行报告生成计划，传递业务场景、分析主题和大纲
            report_results = self._execute_report_plan(
                df, 
                analysis_results, 
                visualization_results, 
                report_plan,
                report_title,
                report_subtitle,
                business_context=business_context,
                analysis_themes=analysis_themes,
                analysis_outline=analysis_outline
            )
            return report_results
        except Exception as e:
            print(f"生成报告时出错: {str(e)}")
            # 如果Agent调用失败，执行基本的报告生成
            return self._basic_report(df, analysis_results, visualization_results, report_title)
    
    def _get_data_overview(self, df: pd.DataFrame) -> str:
        """获取数据概览信息"""
        overview = f"行数: {df.shape[0]}, 列数: {df.shape[1]}\n"
        overview += f"列名: {', '.join(df.columns.tolist())}\n"
        overview += "数据类型:\n"
        
        for col in df.columns:
            dtype = df[col].dtype
            missing = df[col].isna().sum()
            missing_pct = (missing / len(df)) * 100
            overview += f"- {col}: {dtype}, 缺失值: {missing} ({missing_pct:.2f}%)\n"
                
        return overview
    
    def _format_analysis_results(self, analysis_results: Dict[str, Any]) -> str:
        """格式化分析结果，用于提示"""
        formatted = ""
        
        if "descriptive" in analysis_results:
            formatted += "描述性统计分析结果可用\n"
            
        if "correlation" in analysis_results:
            formatted += "相关性分析结果可用\n"
            if "high_correlation_pairs" in analysis_results["correlation"]:
                pairs = analysis_results["correlation"]["high_correlation_pairs"]
                if pairs:
                    formatted += f"发现{len(pairs)}对高相关性特征\n"
        
        if "distribution" in analysis_results:
            formatted += "分布分析结果可用\n"
        
        if "anova" in analysis_results:
            formatted += "方差分析结果可用\n"
        
        if "time_series" in analysis_results:
            formatted += "时间序列分析结果可用\n"
        
        if "clustering" in analysis_results:
            formatted += "聚类分析结果可用\n"
        
        return formatted
    
    def _format_visualization_results(self, visualization_results: Dict[str, Any]) -> str:
        """格式化可视化结果，用于提示"""
        formatted = ""
        
        if "histograms" in visualization_results:
            formatted += f"直方图: {len(visualization_results['histograms'])}个\n"
        
        if "box_plots" in visualization_results:
            formatted += f"箱线图: {len(visualization_results['box_plots'])}个\n"
        
        if "correlation_heatmap" in visualization_results:
            formatted += "相关性热力图可用\n"
        
        if "category_plots" in visualization_results:
            formatted += f"类别计数图: {len(visualization_results['category_plots'])}个\n"
        
        if "scatter_matrix" in visualization_results:
            formatted += "散点图矩阵可用\n"
        
        if "time_series_plots" in visualization_results:
            formatted += f"时间序列图: {len(visualization_results['time_series_plots'])}个\n"
        
        return formatted
    
    def _execute_report_plan(self, df: pd.DataFrame, analysis_results: Dict[str, Any], 
                            visualization_results: Dict[str, Any], report_plan: Any,
                            report_title: str, report_subtitle: str = "",
                            business_context: Optional[Dict[str, Any]] = None,
                            analysis_themes: Optional[List[Dict[str, Any]]] = None,
                            analysis_outline: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """执行报告生成计划"""
        results = {}
        
        # 生成摘要，融合分析主题和大纲
        summary = self._generate_summary_impl(df, analysis_results, visualization_results, 
                                          business_context, analysis_themes, analysis_outline)
        results["summary"] = summary
        
        # 生成引言与背景
        introduction = self._generate_introduction_impl(df, analysis_results, business_context, 
                                                analysis_themes, analysis_outline)
        results["introduction"] = introduction
        
        # 生成数据与方法部分
        data_methods = self._generate_data_methods_impl(df, analysis_results, business_context)
        results["data_methods"] = data_methods
        
        # 生成分析结果呈现部分
        analysis_results_presentation = self._generate_analysis_results_impl(df, analysis_results, visualization_results,
                                                                      business_context, analysis_themes, analysis_outline)
        results["analysis_results_presentation"] = analysis_results_presentation
        
        # 生成结论
        conclusions = self._generate_conclusions_impl(df, analysis_results, business_context, 
                                                analysis_themes, analysis_outline)
        results["conclusions"] = conclusions
        
        # 生成建议，基于分析主题和大纲
        recommendations = self._generate_recommendations_impl(df, analysis_results, visualization_results,
                                                          business_context, analysis_themes, analysis_outline)
        results["recommendations"] = recommendations
        
        # 生成Word报告，整合以上所有内容
        word_report_path = self._generate_word_report_impl(
            df, analysis_results, visualization_results, 
            report_title, report_subtitle, summary, introduction, data_methods,
            analysis_results_presentation, conclusions, recommendations, 
            business_context, analysis_themes, analysis_outline
        )
        
        if word_report_path:
            results["word_report_path"] = word_report_path
        
        return results
    
    def _basic_report(self, df: pd.DataFrame, analysis_results: Dict[str, Any], 
                     visualization_results: Dict[str, Any], report_title: str = "数据分析报告") -> Dict[str, Any]:
        """生成基本的报告"""
        # 与_execute_report_plan相同，但不依赖于Agent的计划
        return self._execute_report_plan(df, analysis_results, visualization_results, None, 
                                       report_title, "数据分析与洞察")
    
    # 工具函数实现
    def _generate_word_report(self, query: str) -> str:
        """生成Word报告工具函数"""
        return "请提供数据框、分析结果和可视化结果，以生成完整的Word报告。"
    
    def _generate_summary(self, query: str) -> str:
        """生成摘要工具函数"""
        return "请提供数据框、分析结果和可视化结果，以生成分析摘要。"
    
    def _generate_recommendations(self, query: str) -> str:
        """生成建议工具函数"""
        return "请提供数据框、分析结果和可视化结果，以生成业务建议。"
    
    def _generate_introduction(self, query: str) -> str:
        """生成引言与背景工具函数"""
        return "请提供数据框、分析结果和业务场景，以生成引言与背景。"
    
    # 实际实现函数
    def _generate_summary_impl(self, df: pd.DataFrame, analysis_results: Dict[str, Any], 
                               visualization_results: Dict[str, Any],
                               business_context: Optional[Dict[str, Any]] = None,
                               analysis_themes: Optional[List[Dict[str, Any]]] = None,
                               analysis_outline: Optional[Dict[str, Any]] = None) -> str:
        """生成数据分析摘要"""
        # 添加业务场景信息
        context_info = ""
        if business_context and "context" in business_context:
            context_info = f"业务场景: {business_context['context']}\n"
            if "explanation" in business_context and business_context["explanation"]:
                context_info += f"场景解释: {business_context['explanation']}\n"
        
        # 添加分析主题信息
        themes_info = ""
        if analysis_themes:
            themes_info = "分析主题:\n"
            for theme in analysis_themes:
                themes_info += f"- {theme['name']}: {theme['description']}\n"
                if "key_questions" in theme and theme["key_questions"]:
                    themes_info += f"  关键问题: {', '.join(theme['key_questions'])}\n"
        
        # 添加分析大纲信息
        outline_info = ""
        if analysis_outline and "outline" in analysis_outline:
            outline_info = "分析大纲结构:\n"
            for section, items in analysis_outline["outline"].items():
                outline_info += f"- {section}\n"
        
        # 构建提示
        prompt_template = PromptTemplate(
            input_variables=["data_info", "analysis_results", "visualization_results", "context_info", "themes_info", "outline_info"],
            template="""
            请为以下数据分析结果生成一个专业的执行摘要。
            
            业务场景信息：
            {context_info}
            
            分析主题：
            {themes_info}
            
            分析大纲：
            {outline_info}
            
            数据集信息：
            {data_info}
            
            分析结果：
            {analysis_results}
            
            可视化结果：
            {visualization_results}
            
            请生成一个300-500字的执行摘要，包括：
            1. 数据集的基本情况和业务背景
            2. 分析主题和目标
            3. 主要分析发现
            4. 关键趋势和模式
            5. 最重要的见解
            
            摘要应该是专业的、基于事实的，并突出最有价值的信息，确保摘要与业务场景和分析主题高度相关。
            摘要不需要包含"摘要"、"执行摘要"等标题，直接开始内容即可。
            """
        )
        
        # 创建LLM链
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        
        # 准备输入
        data_info = self._get_data_overview(df)
        analysis_info = self._format_analysis_results(analysis_results)
        visualization_info = self._format_visualization_results(visualization_results)
        
        # 运行链
        try:
            summary = chain.invoke({
                "data_info": data_info, 
                "analysis_results": analysis_info, 
                "visualization_results": visualization_info,
                "context_info": context_info,
                "themes_info": themes_info,
                "outline_info": outline_info
            })
            return summary["text"]
        except Exception as e:
            print(f"生成摘要时出错: {str(e)}")
            return "无法生成摘要。请查看详细的分析结果。"
    
    def _generate_introduction_impl(self, df: pd.DataFrame, analysis_results: Optional[Dict[str, Any]] = None,
                                   business_context: Optional[Dict[str, Any]] = None,
                                   analysis_themes: Optional[List[Dict[str, Any]]] = None,
                                   analysis_outline: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """生成引言与背景"""
        # 准备业务场景信息
        context_info = ""
        if business_context and "context" in business_context:
            context_info = f"业务场景: {business_context['context']}\n"
            if "explanation" in business_context and business_context["explanation"]:
                context_info += f"场景解释: {business_context['explanation']}\n"
            if "assumptions" in business_context and business_context["assumptions"]:
                context_info += "业务假设:\n"
                for assumption in business_context["assumptions"]:
                    context_info += f"- {assumption}\n"
        
        # 准备分析主题信息
        themes_info = ""
        if analysis_themes:
            themes_info = "分析主题:\n"
            for theme in analysis_themes:
                themes_info += f"- {theme['name']}: {theme['description']}\n"
                if "key_questions" in theme and theme["key_questions"]:
                    themes_info += f"  关键问题: {', '.join(theme['key_questions'])}\n"
        
        # 构建提示
        prompt_template = PromptTemplate(
            input_variables=["data_info", "context_info", "themes_info"],
            template="""
            请为以下数据分析项目生成"引言与背景"部分，包含分析背景和分析目标两个子部分。
            
            业务场景信息：
            {context_info}
            
            分析主题：
            {themes_info}
            
            数据集信息：
            {data_info}
            
            请按照以下格式生成内容：
            
            1. 分析背景（100-150字）：
               - 描述业务需求来源（如管理层决策需求/运营优化需求）
               - 简述相关行业趋势或市场环境（可加入适当的第三方数据支撑）
               - 说明进行此次分析的业务背景和重要性
            
            2. 分析目标（100-150字）：
               - 明确列出待解决的3-5个核心问题
               - 定义成功指标（如"确定影响客户留存的关键因素"）
               - 说明分析结果预期如何支持业务决策
            
            注意：内容应该专业、简洁，总字数控制在200-300字之间。不要包含"引言与背景"等标题，直接输出内容，我会在最终报告中添加适当的标题。使用专业术语，但确保通俗易懂。
            """
        )
        
        # 创建LLM链
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        
        # 准备输入
        data_info = self._get_data_overview(df)
        
        # 运行链
        try:
            result = chain.invoke({
                "data_info": data_info,
                "context_info": context_info,
                "themes_info": themes_info
            })
            
            # 解析结果，分离背景和目标
            background = ""
            goals = ""
            
            text = result["text"]
            parts = text.split("\n\n")
            
            if len(parts) >= 2:
                for part in parts:
                    if "分析背景" in part or "背景" in part:
                        background = part.replace("分析背景：", "").replace("1. 分析背景：", "").strip()
                    elif "分析目标" in part or "目标" in part:
                        goals = part.replace("分析目标：", "").replace("2. 分析目标：", "").strip()
                
                # 如果未能识别出背景和目标，则使用前两段
                if not background and len(parts) > 0:
                    background = parts[0]
                if not goals and len(parts) > 1:
                    goals = parts[1]
            else:
                # 如果文本未按预期分段，则简单地取前半部分作为背景，后半部分作为目标
                half_point = len(text) // 2
                background = text[:half_point].strip()
                goals = text[half_point:].strip()
            
            return {
                "background": background,
                "goals": goals,
                "full_text": text
            }
        except Exception as e:
            print(f"生成引言与背景时出错: {str(e)}")
            return {
                "background": "无法生成分析背景。",
                "goals": "无法生成分析目标。",
                "full_text": "无法生成引言与背景内容。请查看详细的分析结果。"
            }
    
    def _generate_data_methods_impl(self, df: pd.DataFrame, analysis_results: Dict[str, Any],
                                    business_context: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
        """生成数据与方法部分"""
        # 准备业务场景信息
        context_info = ""
        if business_context and "context" in business_context:
            context_info = f"业务场景: {business_context['context']}\n"
        
        # 构建提示
        prompt_template = PromptTemplate(
            input_variables=["data_info", "analysis_results", "context_info"],
            template="""
            请为以下数据分析项目生成"数据与方法"部分，包含数据概况和分析方法两个子部分。
            
            业务场景信息：
            {context_info}
            
            数据集信息：
            {data_info}
            
            分析结果：
            {analysis_results}
            
            请按照以下格式生成内容：
            
            1. 数据概况：
               - 详细描述数据中包含了哪些重要字段（如"销售额、利润、库存"等）
               - 评估数据质量（如完整性、准确性、一致性等）
               - 说明数据的时间跨度
               - 解释通过重要字段能够进行哪些分析（如"按时间维度分析销售额"）
            
            2. 分析方法：
               - 详细说明使用了哪些分析方法（如描述性统计、相关性分析、时间序列分析等）
               - 解释使用这些分析方法的理由和适用性
               - 阐述这些方法如何帮助解答核心问题
               - 提及使用的可视化技术及其作用
            
            注意：内容应该专业、具体，突出数据特点和分析方法的选择理由。不要包含"数据与方法"等标题，直接输出内容，我会在最终报告中添加适当的标题。
            """
        )
        
        # 创建LLM链
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        
        # 准备输入
        data_info = self._get_data_overview(df)
        analysis_info = self._format_analysis_results(analysis_results)
        
        # 运行链
        try:
            result = chain.invoke({
                "data_info": data_info,
                "analysis_results": analysis_info,
                "context_info": context_info
            })
            
            # 解析结果，分离数据概况和分析方法
            data_overview = ""
            methods = ""
            
            text = result["text"]
            parts = text.split("\n\n")
            
            if len(parts) >= 2:
                for part in parts:
                    if "数据概况" in part or "概况" in part:
                        data_overview = part.replace("数据概况：", "").replace("1. 数据概况：", "").strip()
                    elif "分析方法" in part or "方法" in part:
                        methods = part.replace("分析方法：", "").replace("2. 分析方法：", "").strip()
                
                # 如果未能识别出数据概况和分析方法，则使用前两段
                if not data_overview and len(parts) > 0:
                    data_overview = parts[0]
                if not methods and len(parts) > 1:
                    methods = parts[1]
            else:
                # 如果文本未按预期分段，则简单地取前半部分作为数据概况，后半部分作为分析方法
                half_point = len(text) // 2
                data_overview = text[:half_point].strip()
                methods = text[half_point:].strip()
            
            return {
                "data_overview": data_overview,
                "methods": methods,
                "full_text": text
            }
        except Exception as e:
            print(f"生成数据与方法时出错: {str(e)}")
            return {
                "data_overview": "无法生成数据概况。",
                "methods": "无法生成分析方法。",
                "full_text": "无法生成数据与方法内容。请查看详细的分析结果。"
            }
    
    def _generate_analysis_results_impl(self, df: pd.DataFrame, analysis_results: Dict[str, Any], 
                                    visualization_results: Dict[str, Any], business_context: Optional[Dict[str, Any]] = None, 
                                    analysis_themes: Optional[List[Dict[str, Any]]] = None, 
                                    analysis_outline: Optional[Dict[str, Any]] = None) -> str:
        """生成分析结果呈现部分"""
        # 准备业务场景信息
        context_info = ""
        if business_context and "context" in business_context:
            context_info = f"业务场景: {business_context['context']}\n"
        
        # 准备分析主题信息
        themes_info = ""
        if analysis_themes:
            themes_info = "分析主题:\n"
            for theme in analysis_themes:
                themes_info += f"- {theme['name']}: {theme['description']}\n"
        
        # 准备分析大纲信息
        outline_info = ""
        if analysis_outline and "outline" in analysis_outline:
            outline_info = "分析大纲结构:\n"
            for section, items in analysis_outline["outline"].items():
                outline_info += f"- {section}\n"
        
        # 准备可视化图表信息
        vis_info = "可用的可视化图表:\n"
        chart_count = 0
        
        if "histograms" in visualization_results:
            vis_info += f"- 直方图: {len(visualization_results['histograms'])}个\n"
            for col in visualization_results['histograms'].keys():
                vis_info += f"  * {col} 分布直方图\n"
            chart_count += len(visualization_results['histograms'])
        
        if "box_plots" in visualization_results:
            vis_info += f"- 箱线图: {len(visualization_results['box_plots'])}个\n"
            for col in visualization_results['box_plots'].keys():
                vis_info += f"  * {col} 箱线图\n"
            chart_count += len(visualization_results['box_plots'])
        
        if "correlation_heatmap" in visualization_results:
            vis_info += "- 相关性热力图: 1个\n"
            chart_count += 1
        
        if "category_plots" in visualization_results:
            vis_info += f"- 类别分布图: {len(visualization_results['category_plots'])}个\n"
            for col in visualization_results['category_plots'].keys():
                vis_info += f"  * {col} 类别分布图\n"
            chart_count += len(visualization_results['category_plots'])
        
        if "scatter_matrix" in visualization_results:
            vis_info += "- 散点图矩阵: 1个\n"
            chart_count += 1
        
        if "time_series_plots" in visualization_results:
            vis_info += f"- 时间序列图: {len(visualization_results['time_series_plots'])}个\n"
            for key in visualization_results['time_series_plots'].keys():
                cols = key.split("_")
                if len(cols) >= 2:
                    date_col = cols[0]
                    value_col = "_".join(cols[1:])
                    vis_info += f"  * {value_col} 随 {date_col} 的变化趋势图\n"
                else:
                    vis_info += f"  * {key} 时间序列图\n"
            chart_count += len(visualization_results['time_series_plots'])
        
        if chart_count == 0:
            vis_info += "- 无可用图表\n"

        # 准备分析结果信息
        analysis_info = self._format_analysis_results(analysis_results)
        visualization_info = self._format_visualization_results(visualization_results)

        # 构建提示
        prompt_template = PromptTemplate(
            input_variables=["data_info", "analysis_info", "visualization_info", "context_info", "themes_info", "outline_info", "vis_info"],
            template="""
            请为以下数据分析项目生成"分析结果呈现"部分，包含核心指标仪表盘、专题分析模块和关键发现标注。

            数据集信息：
            {data_info}

            分析结果：
            {analysis_info}

            可视化结果：
            {visualization_info}

            业务场景信息：
            {context_info}

            分析主题：
            {themes_info}

            分析大纲：
            {outline_info}
            
            可用的图表：
            {vis_info}

            请按照以下格式生成内容：

            1. 核心指标仪表盘：
               - 包含5-7个关键指标的仪表盘，简要说明每个指标的意义

            2. 专题分析模块：
               - 针对1-3个专题，提供详细的分析结果和图表解读
               - 每个专题模块应引用1-2个相关图表（从可用图表列表中选择）
               - 为每个引用的图表提供专业的解读，包括数据趋势、异常点和业务含义

            3. 关键发现标注：
               - 标注异常数据点和特殊事件，并提供解释

            注意：内容应该专业、清晰，突出最重要的发现，并紧密结合业务场景和分析主题。
            """
        )

        # 创建LLM链
        chain = LLMChain(llm=self.llm, prompt=prompt_template)

        # 准备输入
        data_info = self._get_data_overview(df)

        # 运行链
        try:
            result = chain.invoke({
                "data_info": data_info,
                "analysis_info": analysis_info,
                "visualization_info": visualization_info,
                "context_info": context_info,
                "themes_info": themes_info,
                "outline_info": outline_info,
                "vis_info": vis_info
            })
            return result["text"]
        except Exception as e:
            print(f"生成分析结果呈现时出错: {str(e)}")
            return "无法生成分析结果呈现内容。请查看详细的分析结果。"
    
    def _generate_conclusions_impl(self, df: pd.DataFrame, analysis_results: Dict[str, Any], 
                                    business_context: Optional[Dict[str, Any]] = None, 
                                    analysis_themes: Optional[List[Dict[str, Any]]] = None, 
                                    analysis_outline: Optional[Dict[str, Any]] = None) -> str:
        """生成结论与建议部分"""
        # 准备分析结果信息
        analysis_info = self._format_analysis_results(analysis_results)

        # 构建提示
        prompt_template = PromptTemplate(
            input_variables=["data_info", "analysis_info", "context_info", "themes_info", "outline_info"],
            template="""
            请为以下数据分析项目生成"结论与建议"部分，包含核心结论和可行性建议。

            数据集信息：
            {data_info}

            分析结果：
            {analysis_info}

            业务场景信息：
            {context_info}

            分析主题：
            {themes_info}

            分析大纲：
            {outline_info}

            请按照以下格式生成内容：

            1. 核心结论：
               - 总结分析中发现的最重要的模式、趋势和异常
               - 解释这些发现对业务的意义

            2. 可行性建议：
               - 提供短期行动建议（如优化流程、调整策略）
               - 提供长期规划建议（如投资方向、战略调整）
               - 提供创新方向建议（如新产品、新市场）

            注意：内容应该专业、清晰，突出最重要的发现，并紧密结合业务场景和分析主题。
            """
        )

        # 创建LLM链
        chain = LLMChain(llm=self.llm, prompt=prompt_template)

        # 准备输入
        data_info = self._get_data_overview(df)
        context_info = ""  # 可根据业务场景补充
        themes_info = ""  # 可根据分析主题补充
        outline_info = ""  # 可根据分析大纲补充

        # 运行链
        try:
            result = chain.invoke({
                "data_info": data_info,
                "analysis_info": analysis_info,
                "context_info": context_info,
                "themes_info": themes_info,
                "outline_info": outline_info
            })
            return result["text"]
        except Exception as e:
            print(f"生成结论与建议时出错: {str(e)}")
            return "无法生成结论与建议内容。请查看详细的分析结果。"
    
    def _generate_recommendations_impl(self, df: pd.DataFrame, analysis_results: Dict[str, Any], 
                                       visualization_results: Dict[str, Any], business_context: Optional[Dict[str, Any]] = None, 
                                       analysis_themes: Optional[List[Dict[str, Any]]] = None, 
                                       analysis_outline: Optional[Dict[str, Any]] = None) -> str:
        """生成业务建议部分"""
        # 准备分析结果信息
        analysis_info = self._format_analysis_results(analysis_results)
        visualization_info = self._format_visualization_results(visualization_results)

        # 构建提示
        prompt_template = PromptTemplate(
            input_variables=["data_info", "analysis_info", "visualization_info", "context_info", "themes_info", "outline_info"],
            template="""
            请为以下数据分析项目生成"业务建议"部分，包含短期行动建议、长期规划建议和创新方向建议。

            数据集信息：
            {data_info}

            分析结果：
            {analysis_info}

            可视化结果：
            {visualization_info}

            业务场景信息：
            {context_info}

            分析主题：
            {themes_info}

            分析大纲：
            {outline_info}

            请按照以下格式生成内容：

            1. 短期行动建议：
               - 提供具体的、可立即执行的优化措施

            2. 长期规划建议：
               - 提供战略性规划建议，支持业务的长期发展

            3. 创新方向建议：
               - 提供创新性建议，如新产品、新市场或新技术方向

            注意：内容应该专业、清晰，突出最重要的建议，并紧密结合业务场景和分析主题。
            """
        )

        # 创建LLM链
        chain = LLMChain(llm=self.llm, prompt=prompt_template)

        # 准备输入
        data_info = self._get_data_overview(df)
        context_info = ""  # 可根据业务场景补充
        themes_info = ""  # 可根据分析主题补充
        outline_info = ""  # 可根据分析大纲补充

        # 运行链
        try:
            result = chain.invoke({
                "data_info": data_info,
                "analysis_info": analysis_info,
                "visualization_info": visualization_info,
                "context_info": context_info,
                "themes_info": themes_info,
                "outline_info": outline_info
            })
            return result["text"]
        except Exception as e:
            print(f"生成业务建议时出错: {str(e)}")
            return "无法生成业务建议内容。请查看详细的分析结果。"

    def _generate_word_report_impl(self, df: pd.DataFrame, analysis_results: Dict[str, Any], 
                                   visualization_results: Dict[str, Any], report_title: str, 
                                   report_subtitle: str, summary: str, introduction: Dict[str, str], 
                                   data_methods: Dict[str, str], analysis_results_presentation: str, 
                                   conclusions: str, recommendations: str, 
                                   business_context: Optional[Dict[str, Any]] = None, 
                                   analysis_themes: Optional[List[Dict[str, Any]]] = None, 
                                   analysis_outline: Optional[Dict[str, Any]] = None) -> str:
        """生成完整的Word报告"""
        # 创建Word文档
        document = Document()

        # 添加标题
        document.add_heading(report_title, level=1)
        if report_subtitle:
            document.add_heading(report_subtitle, level=2)

        # 添加摘要
        document.add_heading("摘要", level=2)
        document.add_paragraph(summary)

        # 添加引言与背景
        document.add_heading("引言与背景", level=2)
        document.add_heading("分析背景", level=3)
        document.add_paragraph(introduction.get("background", ""))
        document.add_heading("分析目标", level=3)
        document.add_paragraph(introduction.get("goals", ""))

        # 添加数据与方法
        document.add_heading("数据与方法", level=2)
        document.add_heading("数据概况", level=3)
        document.add_paragraph(data_methods.get("data_overview", ""))
        document.add_heading("分析方法", level=3)
        document.add_paragraph(data_methods.get("methods", ""))

        # 添加分析结果呈现
        document.add_heading("分析结果呈现", level=2)
        document.add_paragraph(analysis_results_presentation)

        # 添加图表和针对图表数据的解读
        document.add_heading("图表与数据解读", level=3)
        document.add_paragraph("以下是针对数据分析生成的图表及其解读：")

        # 处理不同类型的可视化图表
        chart_added = False
        
        # 添加直方图
        if "histograms" in visualization_results:
            document.add_heading("数据分布直方图", level=4)
            document.add_paragraph("直方图展示了数值型变量的分布情况，帮助识别数据的集中趋势、离散程度和异常值。")
            for col, path in visualization_results["histograms"].items():
                try:
                    document.add_heading(f"{col} 分布直方图", level=5)
                    document.add_paragraph(f"该图展示了 {col} 的数值分布情况，可用于分析数据的集中趋势和离散程度。")
                    document.add_picture(path, width=Inches(5))
                    chart_added = True
                except Exception as e:
                    document.add_paragraph(f"无法加载 {col} 直方图: {str(e)}")
        
        # 添加箱线图
        if "box_plots" in visualization_results:
            document.add_heading("数据分布箱线图", level=4)
            document.add_paragraph("箱线图展示了数据的四分位数分布和异常值，有助于理解数据的分散程度和偏斜情况。")
            for col, path in visualization_results["box_plots"].items():
                try:
                    document.add_heading(f"{col} 箱线图", level=5)
                    document.add_paragraph(f"该图展示了 {col} 的四分位数分布，中线表示中位数，箱体表示25%至75%的数据范围，离群点表示潜在异常值。")
                    document.add_picture(path, width=Inches(5))
                    chart_added = True
                except Exception as e:
                    document.add_paragraph(f"无法加载 {col} 箱线图: {str(e)}")
        
        # 添加相关性热力图
        if "correlation_heatmap" in visualization_results:
            document.add_heading("变量相关性热力图", level=4)
            document.add_paragraph("热力图展示了各变量之间的相关性强度，颜色越深表示相关性越强，有助于识别变量间的关联模式。")
            try:
                document.add_picture(visualization_results["correlation_heatmap"], width=Inches(6))
                chart_added = True
            except Exception as e:
                document.add_paragraph(f"无法加载相关性热力图: {str(e)}")
        
        # 添加类别计数图
        if "category_plots" in visualization_results:
            document.add_heading("类别分布图", level=4)
            document.add_paragraph("类别分布图展示了分类变量中各类别的频次分布，有助于理解数据的构成比例。")
            for col, path in visualization_results["category_plots"].items():
                try:
                    document.add_heading(f"{col} 类别分布", level=5)
                    document.add_paragraph(f"该图展示了 {col} 中各类别的分布情况，可用于分析数据的组成结构和主要类别。")
                    document.add_picture(path, width=Inches(5))
                    chart_added = True
                except Exception as e:
                    document.add_paragraph(f"无法加载 {col} 类别分布图: {str(e)}")
        
        # 添加散点图矩阵
        if "scatter_matrix" in visualization_results:
            document.add_heading("变量关系散点图矩阵", level=4)
            document.add_paragraph("散点图矩阵展示了多个数值变量之间的两两关系，有助于发现变量间的相关模式和聚类趋势。")
            try:
                document.add_picture(visualization_results["scatter_matrix"], width=Inches(6))
                chart_added = True
            except Exception as e:
                document.add_paragraph(f"无法加载散点图矩阵: {str(e)}")
        
        # 添加时间序列图
        if "time_series_plots" in visualization_results:
            document.add_heading("时间序列趋势图", level=4)
            document.add_paragraph("时间序列图展示了数据随时间的变化趋势，有助于识别季节性模式、周期性变化和长期趋势。")
            for key, path in visualization_results["time_series_plots"].items():
                try:
                    cols = key.split("_")
                    if len(cols) >= 2:
                        date_col = cols[0]
                        value_col = "_".join(cols[1:])
                        document.add_heading(f"{value_col} 随 {date_col} 的变化趋势", level=5)
                        document.add_paragraph(f"该图展示了 {value_col} 随 {date_col} 的变化趋势，可用于分析时间相关的模式和预测未来走势。")
                    else:
                        document.add_heading(f"{key} 时间序列图", level=5)
                    document.add_picture(path, width=Inches(6))
                    chart_added = True
                except Exception as e:
                    document.add_paragraph(f"无法加载时间序列图 {key}: {str(e)}")
        
        # 如果没有添加任何图表，显示提示信息
        if not chart_added:
            document.add_paragraph("未找到可用的可视化图表。请检查可视化结果或生成新的图表。")
            # 尝试处理旧格式的图表数据
            for chart in visualization_results.get("charts", []):
                document.add_heading(chart.get("title", "图表"), level=4)
                document.add_paragraph(chart.get("description", ""))
                if "path" in chart:
                    try:
                        document.add_picture(chart["path"], width=Inches(5))
                    except Exception as e:
                        document.add_paragraph(f"无法加载图表: {str(e)}")

        # 添加联合业务场景的分析
        document.add_heading("联合业务场景的分析", level=3)
        document.add_paragraph("以下是结合业务场景的多维度分析结果：")
        for analysis in analysis_results.get("business_context_analysis", []):
            document.add_heading(analysis.get("title", "分析"), level=4)
            document.add_paragraph(analysis.get("content", ""))

        # 添加结论与建议
        document.add_heading("结论与建议", level=2)
        document.add_heading("核心结论", level=3)
        document.add_paragraph(conclusions)
        document.add_heading("业务建议", level=3)
        document.add_paragraph(recommendations)

        # 保存文档
        output_path = os.path.join(self.output_dir, f"{report_title}.docx")
        print(f"报告保存路径: {output_path}")
        try:
            document.save(output_path)
            print(f"报告已成功保存到: {output_path}")
            return output_path
        except Exception as e:
            print(f"保存Word报告时出错: {str(e)}")
            return ""

